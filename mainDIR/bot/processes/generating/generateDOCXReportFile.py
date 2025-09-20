import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mainDIR.bot.processes.classesForBot import RequesterData
from businessLogic.classForParsers import CompiledData


async def generateDOCXReport():

    fullname = RequesterData.fullname
    region = RequesterData.region
    birthdate = RequesterData.birthdate
    passport = RequesterData.passport
    passportDate = RequesterData.passportDate

    inn = CompiledData.inn
    fns = CompiledData.fns
    ter = CompiledData.ter
    civ = CompiledData.civ
    bank = CompiledData.bank
    iss = CompiledData.iss

    document = Document()

    # Стиль для заголовков
    heading_style = document.styles['Heading 1']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)

    # Стиль для обычного текста
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(12)

    # Заголовок документа
    document.add_heading('Отчет', level=1)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    birthdate = datetime.datetime.strptime(birthdate, "%Y-%m-%d")
    passportDate = datetime.datetime.strptime(passportDate, "%Y-%m-%d")
    birthdate = datetime.datetime.strftime(birthdate, "%d.%m.%Y")
    passportDate = datetime.datetime.strftime(passportDate, "%d.%m.%Y")

    # Информация о клиенте
    document.add_heading('Информация о запрашиваемом лице', level=2)
    document.add_paragraph(f"ФИО: {fullname}", style='Normal')
    document.add_paragraph(f"Регион: {region}", style='Normal')
    document.add_paragraph(f"Дата рождения: {birthdate}", style='Normal')
    document.add_paragraph(f"Паспорт: {passport}", style='Normal')
    document.add_paragraph(f"Дата выдачи паспорта: {passportDate}", style='Normal')

    # Результаты проверок
    document.add_heading('Результаты проверок', level=2)

    if inn and CompiledData.inn not in ('Информация об ИНН не найдена.', '', "Нет доступа к ресурсу"):
        document.add_paragraph(f"ИНН: {inn}", style='Normal')
        document.add_paragraph(f"ФНС: {fns}", style='Normal')
        document.add_paragraph(f"База террористов: {ter}", style='Normal')
        document.add_paragraph(f"Госслужба: {civ}", style='Normal')
        document.add_paragraph(f"Банкротство: {bank}", style='Normal')
        document.add_paragraph(f"Исполнительные производства:", style='Normal')
        if CompiledData.iss not in ('Ничего не найдено', 'Сервис недоступен'):
            for row in iss:
                if row:
                    for line in row:
                        document.add_paragraph(line, style='Normal')
        else:
            document.add_paragraph(iss, style='Normal')
    else:
        document.add_paragraph(f"ФНС: {fns}", style='Normal')
        document.add_paragraph(f"База террористов: {ter}", style='Normal')
        document.add_paragraph(f"Госслужба: {civ}", style='Normal')
        document.add_paragraph(f"Исполнительные производства:", style='Normal')
        if CompiledData.iss != "Ничего не найдено":
            for row in iss:
                if row:
                    for line in row:
                        document.add_paragraph(line, style='Normal')
        else:
            document.add_paragraph(iss, style='Normal')

    document.save("reportBOT.docx")
