import asyncio
import json
import time
import traceback
from datetime import datetime
import requests
import aiohttp
import docx
import nest_asyncio
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt

nest_asyncio.apply()

filename = "Отчет 2.docx"
api_url = "http://185.33.237.36:5000"
api_key = "123456789"

COUNT_TRY_INN = 2
COUNT_TRY_ALL = 2
WAITING_TIME_INN = 20
WAITING_TIME_ALL = 200
COUNT_TRY_BD = 1
WAITING_TIME_BD = 30


async def iter_visual_cells(row):
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:  # skip cells pointing to same `<w:tc>` element
            continue
        yield cell
        prior_tc = this_tc


import traceback

TOKEN = "5121625794:AAGPhNUcwCR-vMsI-4y4XW-esKK3xAt8XOc"
CHAT_ID = "587990832"


def send_message(text: str):
    url = f"https://api.telegram.org/bot{TOKEN}/sendMessage"
    payload = {
        "chat_id": CHAT_ID,
        "text": text,
        "parse_mode": "Markdown",
    }
    response = requests.post(url, data=payload)


async def call_parser(data, codes, parser_name, sleep_time, count_try, results):
    for attempt in range(count_try):  # Две попытки (можете увеличить, если необходимо)
        async with aiohttp.ClientSession() as session:
            async with session.get(api_url + f'/{parser_name}', params=data) as response:
                result = await response.text()
                result = result.split(' ')[1]
                codes[parser_name] = result
                await asyncio.sleep(sleep_time)

            async with session.get(api_url + '/getResult', params={'api_key': api_key,
                                                                   'key': codes[parser_name]}) as response:
                result = await response.json(content_type=None)

                try:
                    result = json.loads(result)
                except TypeError as ex:
                    result = {
                        "code": '500'
                    }
                try:
                    result.get("code")
                except Exception as ex:
                    print("error", parser_name, result)
                    result = {
                        "code": '500'
                    }

                if result.get("code") == '200':
                    if parser_name == "poiskpasporta":
                        print("poiskpasporta")
                        print(result["message"])
                    if len(result["message"]) == 0:
                        results[parser_name] = "Данные не найдены"
                    elif result["message"] != "Данные не найдены":
                        results[parser_name] = result["message"]
                        break
                    elif result["message"] == "Данные найдены":
                        results[parser_name] = result["message"]
                        break
                    else:
                        results[parser_name] = result["message"]


                else:
                    results[parser_name] = "Данные не найдены"


async def generatePDFreport(id, surname, name, middlename,
                            region, birthday,
                            passport,
                            passport_date, func="text"):
    birthday = birthday.split('-')
    start_time = time.time()
    birthday = birthday[2] + '.' + birthday[1] + '.' + birthday[0]
    passport_date = passport_date.split('-')
    passport_date = passport_date[2] + '.' + passport_date[1] + '.' + passport_date[0]
    doc = docx.Document(filename)

    data = {
        'api_key': api_key,
        'surname': surname,
        'name': name,
        'middlename': middlename,
        'patronymic': middlename,
        'region': region,
        'docnumber': passport,
        'passport': passport,
        'doctype': "21",
        'series': ' '.join(passport.split(' ')[0:2]),
        'number': passport.split(' ')[2],
        'docdate': passport_date,
        'birthdate': birthday,
        'dateBirthday': birthday,
        'date': birthday,
        'fio': f"{surname} {name} {middlename}",
        'InnFio': f"{surname} {name} {middlename}"
    }
    codes = {}
    results = {}

    text = f"Дата формирования отчёта: {datetime.now().strftime('%d.%m.%Y')}"
    doc.sections[0].first_page_header.tables[0].rows[1].cells[0].text = text
    for row in doc.sections[0].first_page_header.tables[0].rows:
        async for cell in iter_visual_cells(row):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Bookman Old Style'
                    font.size = Pt(9)

    for section in doc.sections:
        section.header.tables[0].rows[1].cells[0].text = text
        for row in section.header.tables[0].rows:
            async for cell in iter_visual_cells(row):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Bookman Old Style'
                        font.size = Pt(9)

    # Получение инн
    await call_parser(data, codes, 'inn', WAITING_TIME_INN, COUNT_TRY_INN, results)
    if results['inn'] != "Данные не обнаружены" and results['inn'] != "Данные в базе отстутсвуют " \
            and results['inn'] != "Данные ещё обрабатываются" and results['inn'] != "Данные не найдены":
        data['inn'] = results['inn']['ИНН']

    tasks = [
        call_parser(data, codes, 'declarator', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # Браузер
        call_parser(data, codes, 'gossluzhba', WAITING_TIME_BD, COUNT_TRY_BD, results),  # BD
        call_parser(data, codes, 'criminal', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # requests
        call_parser(data, codes, 'interpol', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # requests
        # call_parser(data, codes, 'rosdolgi', WAITING_TIME_ALL, COUNT_TRY_ALL, results), #Браузер
        call_parser(data, codes, 'kad', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # Браузер
        # call_parser(data, codes, 'ras', WAITING_TIME_ALL, COUNT_TRY_ALL, results),
        call_parser(data, codes, 'natoriat', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # Браузер
        call_parser(data, codes, 'reestzalogov', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # Браузер
        call_parser(data, codes, 'syd', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # Браузер
        call_parser(data, codes, 'VK', 30, 7, results),  # api
        call_parser(data, codes, 'OK', 150, 3, results),  # Браузер
        # NEW
        call_parser(data, codes, 'poiskpasporta', WAITING_TIME_BD, COUNT_TRY_BD, results),  # BD
        call_parser(data, codes, 'terrorist', WAITING_TIME_BD, COUNT_TRY_BD, results),  # BD
        call_parser(data, codes, 'chek_diskvalif', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # Браузер
        call_parser(data, codes, 'ycherediteley', WAITING_TIME_BD, COUNT_TRY_BD, results),  # BD
        # NEW
        call_parser(data, codes, 'advokat', WAITING_TIME_BD, COUNT_TRY_BD, results),  # BD
        call_parser(data, codes, 'natoriys', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # BD
        call_parser(data, codes, 'konkursyprv', WAITING_TIME_ALL, COUNT_TRY_ALL, results),  # BD
        call_parser(data, codes, 'inagent', WAITING_TIME_BD, COUNT_TRY_BD, results),  # BD

    ]

    if data.get('inn') != None:
        tasks.append(call_parser(data, codes, 'fns', WAITING_TIME_ALL, COUNT_TRY_ALL, results))  # requests
        tasks.append(call_parser(data, codes, 'checkstatus', WAITING_TIME_ALL, COUNT_TRY_ALL, results)),  # requests
        tasks.append(call_parser(data, codes, 'inndolg', WAITING_TIME_ALL, COUNT_TRY_ALL, results)),  # Браузер
        tasks.append(call_parser(data, codes, 'EGRUL', WAITING_TIME_ALL, COUNT_TRY_ALL, results)),  # Браузер
        tasks.append(call_parser(data, codes, 'bankrot', WAITING_TIME_ALL, COUNT_TRY_ALL, results))  # Браузер
    else:
        results['fns'] = 'Данные не найдены'
        results['checkstatus'] = 'Данные не найдены'
        results['inndolg'] = 'Данные не найдены'
        results['EGRUL'] = 'Данные не найдены'
        results['bankrot'] = 'Данные не найдены'
    # results['advokat'] = 'Данные не найдены'
    # results['natoriys'] = 'Данные не найдены'
    # results['konkursyprv'] = 'Данные не найдены'
    # results['inagent'] = 'Данные не найдены'
    results['rosdolgi'] = 'Данные не найдены'

    parser_loop = asyncio.get_event_loop()
    results_tasks = parser_loop.run_until_complete(asyncio.gather(*tasks))

    data_type_dict = ['ФИО:', 'Дата рождения:', 'Серия, номер паспорта:', 'Дата выдачи:', 'Регион поиска:',
                      'ИНН:', 'Паспорт:', 'Розыск:', 'Экстремистская деятельность и терроризм:',
                      "Налоговые начисления:",
                      'Исполнительные производства:', 'Залог движимого имущества:', 'Наследственное дело:',
                      'Сведения о банкротстве:', 'Суды общей юрисдикции:', 'Арбитражные суды:', 'Статус самозанятого:',
                      'ЕГРИП:', 'ЕГРЮЛ:', 'Сведения о дисквалификации:', 'Реестр массовых учредителей:',
                      'Сведения о декларациях:',
                      'Сведения об утрате доверия:', 'Социальные сети:', "Реестр нотариусов:", "Реестр адвокатов:",
                      "Реестр арбитражных управляющих:", "Иностранные агенты:",
                      ]
    states = ""
    doc.paragraphs[0].text = ""
    run = doc.paragraphs[0].add_run(f"Отчёт № {id}")
    main_style = run.style
    main_style.font.size = Pt(14)
    main_style.font.name = 'Bookman Old Style'
    main_style.font.bold = True

    for parser in results:
        if parser != 'inn':
            if results[parser] == "Данный источник не доступен,при доступности данные будут доотправлены" or \
                    results[parser] == "Данные ещё обрабатываются":
                # print(f"[ERROR] {results[parser]}")
                results[parser]["message"] = "Данные не найдены"
    for table in doc.tables:
        for row in table.rows:
            async for cell in iter_visual_cells(row):
                if states == 'ФИО:' and cell.text not in data_type_dict:
                    cell.text = surname + ' ' + name + ' ' + middlename
                if states == 'Дата рождения:' and cell.text not in data_type_dict:
                    cell.text = birthday
                if states == 'Серия, номер паспорта:' and cell.text not in data_type_dict:
                    cell.text = passport
                if states == 'Дата выдачи:' and cell.text not in data_type_dict:
                    cell.text = passport_date
                if states == 'Регион поиска:' and cell.text not in data_type_dict:
                    cell.text = region

                if states == 'ИНН:' and cell.text not in data_type_dict:
                    inn_data = results["inn"]
                    if not isinstance(inn_data, str):
                        inn_data = inn_data["ИНН"]
                    if inn_data == "Данные не обнаружены" or inn_data == "Данные в базе отстутсвуют ":
                        inn_data = "данных не найдено"
                    cell.text = inn_data
                if states == 'Паспорт:' and cell.text not in data_type_dict:
                    passport_data = results["poiskpasporta"]
                    if passport_data == "Данные  найдены":
                        passport_data = "числится в базе недействительных"
                    if passport_data == "Данные не найдены":
                        passport_data = "данных о недействительности паспорта не найдено"
                    cell.text = passport_data
                if states == 'Розыск:' and cell.text not in data_type_dict:
                    criminal_data = results['criminal']
                    if criminal_data == "Данные не найдены":
                        criminal_data = "сведений о розыске лица по данным ФСИН не найдено"
                    else:
                        criminal_data = "числится в розыске по данным ФСИН"
                    interpol_data = results['interpol']
                    if interpol_data == "Данные не найдены" or interpol_data == "Не числится в Интерполе":
                        interpol_data = "сведений о розыске лица по данным Интерпол не найдено"
                    else:
                        interpol_data = "числится в розыске по данным Интерпол"
                    text = f"{criminal_data}\n{interpol_data}"
                    cell.text = text

                if states == 'Экстремистская деятельность и терроризм:' and cell.text not in data_type_dict:
                    terrorist_data = results["terrorist"]
                    if terrorist_data == "Данные не найдены":
                        terrorist_data = "данных о внесении в реестр лиц, " \
                                         "в отношении которых имеются сведения о причастности к" \
                                         " экстремистской деятельности или терроризму, не найдено "
                    else:
                        terrorist_data = "числится в реестре лиц, " \
                                         "в отношении которых имеются сведения о причастности к " \
                                         "экстремистской деятельности или терроризму"
                    cell.text = terrorist_data

                if states == 'Иностранные агенты:' and cell.text not in data_type_dict:
                    inagent_data = results["inagent"]
                    if inagent_data == "Данные не найдены":
                        inagent_data = "данных о внесении в реестр иностранных агентов не найдено"
                    else:
                        inagent_data = "числится в реестре иностранных агентов"
                    cell.text = inagent_data

                if states == 'Налоговые начисления:' and cell.text not in data_type_dict:
                    inndolg_data = results["inndolg"]
                    if inndolg_data == "Данные не найдены" or inndolg_data == "error":
                        text = "данных о налоговых начислениях не найдено"
                    else:
                        text = 'найдены сведения о налоговых начислениях:\n'
                        if not isinstance(inndolg_data, str):
                            for line in inndolg_data:
                                _ = f"ИНН: {line['ИНН']}\n" \
                                    f"Получатель платежа: {line['Получатель платежа']}\n" \
                                    f"Расчетный счет: {line['Расчетный счет']}\n" \
                                    f"БИК: {line['БИК']}\n" \
                                    f"ОКТМО(ОКАТО): {line['ОКТМО(ОКАТО)']}\n" \
                                    f"КБК: {line['КБК']}\n" \
                                    f"Сумма: {line['Сумма']}\n\n"
                                text = text + _
                    cell.text = text

                if states == 'Исполнительные производства:' and cell.text not in data_type_dict:
                    rosdolgi_data = results['rosdolgi']
                    if rosdolgi_data == "Данные не найдены":
                        result_text = "данных об исполнительных производствах не найдено"
                    else:
                        result_text = 'найдены сведения об исполнительных производствах:\n'
                        if not isinstance(rosdolgi_data, str):
                            for line in rosdolgi_data:
                                text = f"Номер и дата испол. производства: {line['Номер и дата испол. производства']}\n" \
                                       f"Реквизиты испол. документа: {line['Реквизиты испол. документа']}\n" \
                                       f"Предмет исполнения: {line['Предмет исполнения']}\n" \
                                       f"Отдел судебных приставов: {line['Отдел судебных приставов']}\n" \
                                       f"Пристав-исполнитель: {line['Пристав-исполнитель']}\n\n"
                                result_text = result_text + text

                    cell.text = result_text

                if states == "Залог движимого имущества:" and cell.text not in data_type_dict:
                    reestzalogov_data = results['reestzalogov']
                    if reestzalogov_data == "Данные не найдены":
                        result_text = "данных о внесении сведений о лице в реестр уведомлений о " \
                                      "залоге движимого имущества не найдено"
                    else:
                        result_text = 'числится в реестре уведомлений о ' \
                                      'залоге движимого имущества в качестве залогодателя:\n'
                        if not isinstance(reestzalogov_data, str):
                            for line in reestzalogov_data:
                                text = f"Дата регистрации: {line['Дата регистрации']}\n" \
                                       f"Номер уведомления о возникновении залога: {line['Номер уведомления о возникновении залога']}\n" \
                                       f"Имущество: {line['Имущество']}\n" \
                                       f"Залогодатель: {line['Залогодатель']}\n" \
                                       f"Залогодержатель: {line['Залогодержатель']}\n\n"
                                result_text = result_text + text
                    cell.text = result_text

                if states == "Наследственное дело:" and cell.text not in data_type_dict:
                    natoriat_data = results['natoriat']
                    if natoriat_data == "Данные не найдены":
                        result_text = "данных об открытии наследственного дела не найдено"
                    else:
                        result_text = 'найдены сведения об открытии наследственного дела:\n'
                        if not isinstance(natoriat_data, str):
                            for line in natoriat_data:
                                text = f"Дата смерти: {line['Дата смерти']}\n" \
                                       f"Номер наследственного дела: {line['Номер наследственного дела']}\n" \
                                       f"Нотариус: {line['Нотариус']}\n\n"
                                result_text = result_text + text
                    cell.text = result_text

                # НЕ РЕАЛИЗОВАНО. ТЗ_4
                if states == "Сведения о банкротстве:" and cell.text not in data_type_dict:
                    bankrot_data = results['bankrot']
                    if bankrot_data == "Данные не найдены":
                        result_text = "          сведений о банкротстве физического лица не найдено\n "
                    else:
                        result_text = '          найдены сведения о банкротстве физического лица:\n'
                        if not isinstance(bankrot_data, str):
                            text = f"Результат: {bankrot_data['status']}\n" \
                                   f"Номер судебного дела: {bankrot_data['nomer']}\n" \
                                   f"Снилс: {bankrot_data['snils']}\n" \
                                   f"Арбитражный управляющий: {bankrot_data['arbitr']}\n\n"
                            result_text = result_text + text
                    # result_text = "Привет"
                    cell.text = result_text

                # Нужно сделать парсер новый, пока выводим пустым
                result_text = ""
                if states == "Суды общей юрисдикции:" and cell.text not in data_type_dict:
                    syd_data = results['syd']
                    if syd_data == "Данные не найдены":
                        result_text = "данных об участии в делах, рассматриваемых судами общей юрисдикции, не найдено"
                    else:
                        if not isinstance(syd_data, str):
                            for line in syd_data:
                                text = f"Суд: {line['Суд']}\n" \
                                       f"Номер дела: {line['Номер дела']}\n" \
                                       f"Движение дела: {line['Движение дела']}\n" \
                                       f"Участники: {line['Участники']}\n\n"
                                result_text = result_text + text
                    # result_text = "Привет"
                    cell.text = result_text

                if states == "Арбитражные суды:" and cell.text not in data_type_dict:
                    kad_data = results['kad']
                    if kad_data == "Данные не найдены":
                        result_text = "данных об участии в делах, рассматриваемых арбитражными судами, не найдено"
                    else:
                        result_text = ""
                        # result_text = 'найдены сведе об участии в делах, рассматриваемых арбитражными судами:\n'
                        if not isinstance(kad_data, str):
                            for line in kad_data:
                                text = f"Дело: {line['case_number']}, {line['date']}n" \
                                       f"Судья | текущая инстанция: {line['court']}\n" \
                                       f"Истец: {line['plaintiff']}\n" \
                                       f"Ответчик: {line['respondent']}\n\n"
                                result_text = result_text + text
                    cell.text = result_text

                # checkstatus
                if states == "Статус самозанятого:" and cell.text not in data_type_dict:
                    checkstatus_data = results['checkstatus']
                    if checkstatus_data == "Данные не найдены":
                        checkstatus_data = "данных о регистрации в качестве плательщика " \
                                           "налога на профессиональный доход (самозанятого) не найдено"
                    else:
                        checkstatus_data = "найдены сведения о регистрации в качестве плательщика " \
                                           "налога на профессиональный доход (самозанятого)"
                    cell.text = checkstatus_data

                if states == "ЕГРИП:" and cell.text not in data_type_dict:
                    fns_data = results['fns']
                    if fns_data == "Данные не найдены":
                        fns_data = "данных о регистрации в качестве индивидуального предпринимателя не найдено"
                    else:
                        fns_data = "зарегистрирован(а) в качестве индивидуального предпринимателя"
                    cell.text = str(fns_data)

                if states == "ЕГРЮЛ:" and cell.text not in data_type_dict:
                    EGRUL_data = results['EGRUL']
                    if EGRUL_data == "Данные не найдены":
                        result_text = "данных о регистрации в качестве учредителя и/или руководителя организации не найдено"
                    else:
                        result_text = 'зарегистрирован(а) в качестве учредителя и/или руководителя организации:\n\n'
                        # result_text = result_text + '\n'
                        if not isinstance(EGRUL_data, str):
                            for line in EGRUL_data:
                                text = f"Наименование: {line['name']}\n" \
                                       f"ИНН: {line['inn']}\n" \
                                       f"Регион: {line['region']}\n\n"
                                result_text = result_text + text
                    cell.text = result_text

                if states == "Сведения о дисквалификации:" and cell.text not in data_type_dict:
                    chek_diskvalif_data = results['chek_diskvalif']
                    if chek_diskvalif_data == "Данные не найдены":
                        result_text = "данных о внесении в реестр дисквалифицированных лиц ФНС РФ не найдено"
                    else:
                        result_text = 'числится в реестре дисквалифицированных лиц ФНС РФ:\n'
                        if not isinstance(chek_diskvalif_data, str):
                            text = str(chek_diskvalif_data)
                            result_text = result_text + text
                    cell.text = result_text

                if states == 'Реестр массовых учредителей:' and cell.text not in data_type_dict:
                    ycherediteley_data = results['ycherediteley']
                    if ycherediteley_data == "Данные не найдены":
                        result_text = "данных о внесении в реестр массовых учредителей юридических лиц не найдено"
                    else:
                        result_text = 'числится в реестре массовых учредителей юридических лиц:\n'
                        if not isinstance(ycherediteley_data, str):
                            text = str(ycherediteley_data)
                            result_text = result_text + text
                    cell.text = result_text

                if states == 'Реестр нотариусов:' and cell.text not in data_type_dict:
                    natoriys_data = results['natoriys']
                    if natoriys_data == "Данные не найдены":
                        result_text = "данных о внесении в реестр нотариусов не найдено"
                    else:
                        result_text = ""
                        # result_text = 'числится в реестре нотариусов:\n'
                        for line in natoriys_data:
                            text = f"Статус: {line['Статус']}\n" \
                                   f"Адрес: {line['Адрес']}\n" \
                                   f"Телефон: {line['Телефон']}\n"
                        result_text = text

                    cell.text = result_text

                if states == 'Реестр адвокатов:' and cell.text not in data_type_dict:
                    advokat_data = results['advokat']
                    if advokat_data == "Данные не найдены":
                        result_text = "данных о внесении в реестр адвокатов не найдено"
                    else:
                        for line in advokat_data:
                            text = f"Регистрационный номер: {line['Регистрационный номер']}\n" \
                                   f"Субъект РФ: {line['Субъект РФ']}\n" \
                                   f"Номер удостоверения: {line['Номер удостоверения']}\n" \
                                   f"Статус: {line['Статус']}"
                        result_text = text

                    cell.text = result_text

                if states == 'Реестр арбитражных управляющих:' and cell.text not in data_type_dict:
                    konkursyprv_data = results['konkursyprv']
                    if konkursyprv_data == "Данные не найдены":
                        result_text = "данных о внесении в реестр конкурсных управляющих не найдено"
                    else:
                        result_text = ""
                        # result_text = 'данные в реестре конкрусных управляющих найдены:\n'
                        for line in konkursyprv_data:
                            text = f"Рег. номер: {line['Рег. номер']}\n" \
                                   f"Дата регистрации в Росреестре: {line['Дата регистрации в Росреестре']}\n" \
                                   f"СРО: {line['СРО']}\n"
                        result_text = result_text + text

                    cell.text = result_text

                if states == 'Сведения о декларациях:' and cell.text not in data_type_dict:
                    declorator_data = results['declarator']
                    if declorator_data == "Данные не найдены":
                        result_text = "данных о декларировании доходов и имущества не найдено"
                    else:
                        result_text = ""
                        # result_text = 'найдены сведения о декларировании доходов и имущества:\n'
                        if not isinstance(declorator_data, str):
                            for line in declorator_data:
                                text = f"Тип декларации и год дохода: {line['Тип декларации и год дохода']}\n" \
                                       f"Доход, руб.: {line['Доход, руб.']}\n" \
                                       f"Недвижимость, м2: {line['Недвижимость, м2']}\n" \
                                       f"Транспорт, шт.: {line['Транспорт, шт.']}\n\n"
                                result_text = result_text + text
                    cell.text = result_text

                if states == "Сведения об утрате доверия:" and cell.text not in data_type_dict:
                    gossluzhba_data = results['gossluzhba']
                    if gossluzhba_data == "Данные не найдены":
                        result_text = "данных о внесении в реестр лиц, уволенных с " \
                                      "должностей государственной службы в связи с утратой доверия, не найдено"
                    else:
                        result_text = 'числится в реестре лиц, уволенных с должностей государственной' \
                                      ' службы в связи с утратой доверия'

                    cell.text = result_text

                if states == "Вконтакте" and cell.text not in data_type_dict:
                    VK_data = results['VK']
                    if VK_data == "Данные не найдены":
                        result_text = "данных о регистрации в социальных сетях не найдено"
                    else:
                        result_text = '\n'
                        if not isinstance(VK_data, str):
                            d = 0
                            for line in VK_data:
                                if 'ok.ru' not in line['link']:
                                    d = d + 1
                                    text = f"{d}. Имя и фамилия: {line['info']}\n" \
                                           f"Ссылка: {line['link']}\n\n"
                                    result_text = result_text + text
                    cell.text = result_text

                if states == "Одноклассники" and cell.text not in data_type_dict:
                    VK_data = results['OK']
                    if VK_data == "Данные не найдены":
                        result_text = "данных о регистрации в социальных сетях не найдено"
                    else:
                        result_text = '\n'
                        if not isinstance(VK_data, str):
                            d = 0
                            for line in VK_data:
                                if 'vk.com' not in line['link']:
                                    d = d + 1
                                    text = f"{d}. Имя и фамилия: {line['info']}\n" \
                                           f"Ссылка: {line['link']}\n\n"
                                    result_text = result_text + text
                    cell.text = result_text

                if states == 'Дата формирования отчёта:' and cell.text not in data_type_dict:
                    cell.text = cell.text.replace('Дата формирования отчёта:', f"Дата формирования отчёта: "
                                                                               f"{datetime.now()}")

                if states != "":
                    states = ""

                # Общие данные

                if 'ФИО:' in cell.text: states = 'ФИО:'
                if 'Дата рождения:' in cell.text: states = 'Дата рождения:'
                if 'Серия, номер паспорта:' in cell.text: states = 'Серия, номер паспорта:'
                if 'Дата выдачи:' in cell.text: states = 'Дата выдачи:'
                if 'Регион поиска:' in cell.text: states = 'Регион поиска:'
                if 'ИНН:' in cell.text: states = 'ИНН:'

                # Данные правоохранительных органов:
                if 'Паспорт:' in cell.text: states = 'Паспорт:'
                if 'Розыск:' in cell.text: states = 'Розыск:'
                if 'Экстремистская деятельность и терроризм:' in cell.text:
                    states = 'Экстремистская деятельность и терроризм:'

                if 'Иностранные агенты:' in cell.text: states = 'Иностранные агенты:'

                # Имущественные права и обязательства:
                if 'Налоговые начисления:' in cell.text: states = 'Налоговые начисления:'
                if 'Исполнительные производства:' in cell.text: states = 'Исполнительные производства:'
                if 'Залог движимого имущества:' in cell.text: states = 'Залог движимого имущества:'
                if 'Наследственное дело:' in cell.text: states = 'Наследственное дело:'

                if 'Сведения о банкротстве:' in cell.text: states = 'Сведения о банкротстве:'

                # Участие в судебных делах:
                if 'Суды общей юрисдикции:' in cell.text: states = 'Суды общей юрисдикции:'
                if 'Арбитражные суды:' in cell.text: states = 'Арбитражные суды:'

                # Сведения о предпринимательской и иной экономической деятельности:
                if 'Статус самозанятого:' in cell.text: states = 'Статус самозанятого:'
                if 'ЕГРИП:' in cell.text: states = 'ЕГРИП:'
                if 'ЕГРЮЛ:' in cell.text: states = 'ЕГРЮЛ:'
                if 'Сведения о дисквалификации:' in cell.text: states = 'Сведения о дисквалификации:'
                if 'Реестр массовых учредителей:' in cell.text: states = 'Реестр массовых учредителей:'

                if 'Реестр нотариусов:' in cell.text: states = 'Реестр нотариусов:'
                if 'Реестр адвокатов:' in cell.text: states = 'Реестр адвокатов:'
                if 'Реестр арбитражных управляющих:' in cell.text: states = 'Реестр арбитражных управляющих:'

                # Госслужба:
                if 'Сведения о декларациях:' in cell.text: states = 'Сведения о декларациях:'
                if 'Сведения об утрате доверия:' in cell.text: states = 'Сведения об утрате доверия:'

                if 'Вконтакте' in cell.text: states = 'Вконтакте'

                if 'Одноклассники' in cell.text: states = 'Одноклассники'

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Bookman Old Style'

    file_path = r'Отчёты/' + f"Отчёт № {id}.docx"
    doc.save(file_path)
    end_time = time.time()
    execution_time = end_time - start_time
    messege = f"Время формирования отчета № {id} за  {execution_time} вызов из фунции {func}"
    send_message(messege)

# async def main():
#    try:
#        #await generatePDFreport(22, "Хожило", "Валерий", "Анатольевич", "Магаданская область", "1975-07-25",
#        #                        "44 03 176401",
#        #                        "2009-09-22")
#        await generatePDFreport(20, "Грибов", "Роман", "Иванович", "Магаданская область", "1978-08-22",
#                                                                                          "44 08 246344",
#                                                                                          "2008-04-01")
#    except Exception as ex:
#        print(f"!!!!{ex}")
#        traceback.print_exc()
# if __name__ == '__main__':
#    loop = asyncio.get_event_loop()
#    loop.run_until_complete(main())
#    loop.close()

