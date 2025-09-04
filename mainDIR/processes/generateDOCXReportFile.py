
import datetime
from sqlite3 import connect

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


async def generateDOCXReport(
        result: dict,
        ID: int | tuple,
        fullname: str,
        region: str,
        birthdate: datetime.date,
        passport: str,
        passportDate: datetime.date
) -> str:
    if type(ID) is tuple:
        ID = ID[0]
    fullnameForFilename = fullname.replace(' ', '-')
    filename = f"Report-{ID}-{fullnameForFilename}.docx"

    document = Document()

    # Стиль для заголовков
    heading_style = document.styles['Heading 1']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)

    # Стиль для обычного текста
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(12)

    # Заголовок документа
    document.add_heading('Отчет', level=1)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Информация о клиенте
    document.add_heading('Информация о клиенте', level=2)
    document.add_paragraph(f"ФИО: {fullname}", style='Normal')
    document.add_paragraph(f"Регион: {region}", style='Normal')
    document.add_paragraph(f"Дата рождения: {birthdate}", style='Normal')
    document.add_paragraph(f"Паспорт: {passport}", style='Normal')
    document.add_paragraph(f"Дата выдачи паспорта: {passportDate}", style='Normal')

    # Результаты проверок
    document.add_heading('Результаты проверок', level=2)

    if result.get('inn') and result.get('inn') != 'Информация об ИНН не найдена.' and result.get('inn') != '' and result.get('inn') != "Нет доступа к ресурсу":
        document.add_paragraph(f"ИНН: {result.get('inn').get('inn')}", style='Normal')
        document.add_paragraph(f"ФНС: {result.get('fns')}", style='Normal')
        document.add_paragraph(f"База террористов: {result.get('ter')}", style='Normal')
        document.add_paragraph(f"Госслужба: {result.get('civ')}", style='Normal')
        document.add_paragraph(f"Банкротство: {result.get('bank')}", style='Normal')
        document.add_paragraph(f"Исполнительные производства:", style='Normal')
        if result.get('iss') != "Ничего не найдено":
            for row in result.get('iss'):
                if row:
                    for line in row:
                        document.add_paragraph(line, style='Normal')
        else:
            document.add_paragraph("Ничего не найдено", style='Normal')
    else:
        document.add_paragraph(f"ФНС: {result.get('fns')}", style='Normal')
        document.add_paragraph(f"База террористов: {result.get('ter')}", style='Normal')
        document.add_paragraph(f"Госслужба: {result.get('civ')}", style='Normal')
        document.add_paragraph(f"Исполнительные производства:", style='Normal')
        if result.get('iss') != "Ничего не найдено":
            for row in result.get('iss'):
                if row:
                    for line in row:
                        document.add_paragraph(line, style='Normal')
        else:
            document.add_paragraph("Ничего не найдено", style='Normal')

    document.save(filename)
    return filename
