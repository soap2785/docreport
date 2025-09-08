import asyncio
import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mainDIR.processes.bot.classes import RequesterData


async def generateDOCXReport(
        ID: int | tuple,
        CompiledData: dict
) -> str:
    if type(ID) is tuple:
        ID = ID[0]
    fullnameForFilename = RequesterData.fullname.replace(' ', '-')
    filename = f"Report-{ID}-{fullnameForFilename}.docx"

    fullname = RequesterData.fullname
    region = RequesterData.region
    birthdate = RequesterData.birthdate
    passport = RequesterData.passport
    passportDate = RequesterData.passportDate

    document = Document()

    # Стиль для заголовков
    heading_style = document.styles['Heading 1']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)

    # Стиль для обычного текста
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(12)

    # Заголовок документа
    document.add_heading('Отчет', level=1)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    birthdate = datetime.datetime.strftime(birthdate, "%d.%m.%Y")
    passportDate = datetime.datetime.strftime(passportDate, "%d.%m.%Y")

    # Информация о клиенте
    document.add_heading('Информация о запрашиваемом лице', level=2)
    document.add_paragraph(f"ФИО: {fullname}", style='Normal')
    document.add_paragraph(f"Регион: {region}", style='Normal')
    document.add_paragraph(f"Дата рождения: {birthdate}", style='Normal')
    document.add_paragraph(f"Паспорт: {passport}", style='Normal')
    document.add_paragraph(f"Дата выдачи паспорта: {passportDate}", style='Normal')

    # Результаты проверок
    document.add_heading('Результаты проверок', level=2)

    if CompiledData.get('inn') and CompiledData.get('inn') != 'Информация об ИНН не найдена.' and CompiledData.get('inn') != '' and CompiledData.get('inn') != "Нет доступа к ресурсу":
        document.add_paragraph(f"ИНН: {CompiledData.get('inn').get('inn')}", style='Normal')
        document.add_paragraph(f"ФНС: {CompiledData.get('fns')}", style='Normal')
        document.add_paragraph(f"База террористов: {CompiledData.get('ter')}", style='Normal')
        document.add_paragraph(f"Госслужба: {CompiledData.get('civ')}", style='Normal')
        document.add_paragraph(f"Банкротство: {CompiledData.get('bank')}", style='Normal')
        document.add_paragraph(f"Исполнительные производства:", style='Normal')
        if CompiledData.get('iss') != "Ничего не найдено":
            for row in CompiledData.get('iss'):
                if row:
                    for line in row:
                        document.add_paragraph(line, style='Normal')
        else:
            document.add_paragraph("Ничего не найдено", style='Normal')
    else:
        document.add_paragraph(f"ФНС: {CompiledData.get('fns')}", style='Normal')
        document.add_paragraph(f"База террористов: {CompiledData.get('ter')}", style='Normal')
        document.add_paragraph(f"Госслужба: {CompiledData.get('civ')}", style='Normal')
        document.add_paragraph(f"Исполнительные производства:", style='Normal')
        if CompiledData.get('iss') != "Ничего не найдено":
            for row in CompiledData.get('iss'):
                if row:
                    for line in row:
                        document.add_paragraph(line, style='Normal')
        else:
            document.add_paragraph("Ничего не найдено", style='Normal')

    document.save(filename)
    await asyncio.sleep(1)
    return filename
